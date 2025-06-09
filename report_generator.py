from io import BytesIO
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import base64
import numpy as np

class EnhancedReportGenerator:
    """
    Enhanced report generator that supports reader/writer sizing recommendations
    """
    
    def __init__(self):
        self.colors = {
            'writer': '#2196F3',
            'reader': '#9C27B0', 
            'storage': '#4CAF50',
            'backup': '#FF9800',
            'primary': '#1976D2',
            'secondary': '#7B1FA2'
        }
    
    def generate_pdf_report(self, calculator):
        """Generate comprehensive PDF report with reader/writer analysis"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch)
            styles = getSampleStyleSheet()
            elements = []
            
            # Title Page
            title_style = ParagraphStyle(
                name="Title",
                fontSize=24,
                alignment=1,
                spaceAfter=20,
                fontName="Helvetica-Bold",
                textColor=colors.HexColor("#1976D2")
            )
            elements.append(Paragraph("AWS RDS/Aurora Sizing Report", title_style))
            elements.append(Paragraph("Reader/Writer Optimized Recommendations", 
                                    ParagraphStyle(name="Subtitle", fontSize=16, alignment=1, 
                                                 spaceAfter=30, fontName="Helvetica", 
                                                 textColor=colors.HexColor("#666"))))
            
            # Executive Summary
            elements.append(Paragraph("Executive Summary", styles['Heading1']))
            
            # Get deployment configuration
            deployment = calculator.inputs.get("deployment", "Multi-AZ")
            deployment_config = calculator.DEPLOYMENT_OPTIONS.get(deployment, {})
            has_readers = deployment_config.get("has_readers", False)
            
            summary_text = f"""
            <b>Database Engine:</b> {calculator.inputs.get("engine", "N/A")}<br/>
            <b>Region:</b> {calculator.inputs.get("region", "N/A")}<br/>
            <b>Deployment Type:</b> {deployment}<br/>
            <b>Reader Instances:</b> {"Yes (" + str(deployment_config.get("reader_count", 0)) + " per environment)" if has_readers else "No (Single-AZ)"}<br/>
            <b>Workload Pattern:</b> {calculator.inputs.get("workload_pattern", "N/A")}<br/>
            """
            
            if "PROD" in calculator.recommendations:
                prod_rec = calculator.recommendations["PROD"]
                if "error" not in prod_rec:
                    summary_text += f"""
                    <b>Production Writer:</b> {prod_rec.get("writer", {}).get("instance_type", "N/A")}<br/>
                    """
                    if prod_rec.get("readers"):
                        readers = prod_rec["readers"]
                        summary_text += f"""
                        <b>Production Readers:</b> {readers.get("instance_type", "N/A")} x{readers.get("count", 0)}<br/>
                        """
                    summary_text += f"""
                    <b>Total Monthly Cost:</b> ${prod_rec.get("total_cost", 0):,.2f}<br/>
                    """
            
            elements.append(Paragraph(summary_text, styles['Normal']))
            elements.append(Spacer(1, 20))
            
            # Deployment Architecture
            elements.append(Paragraph("Deployment Architecture", styles['Heading2']))
            
            arch_description = self._generate_architecture_description(calculator)
            elements.append(Paragraph(arch_description, styles['Normal']))
            elements.append(Spacer(1, 15))
            
            # Cost Summary Table
            elements.append(Paragraph("Cost Summary by Environment", styles['Heading2']))
            
            cost_data = [["Environment", "Writer Instance", "Writer Cost", "Reader Instance", "Reader Cost", "Storage Cost", "Total Cost"]]
            
            for env in calculator.recommendations:
                rec = calculator.recommendations[env]
                if "error" not in rec:
                    row = [
                        env,
                        rec.get("writer", {}).get("instance_type", "N/A"),
                        f"${rec.get('instance_cost', 0):,.2f}",
                    ]
                    
                    if rec.get("readers"):
                        readers = rec["readers"]
                        row.extend([
                            f"{readers.get('instance_type', 'N/A')} x{readers.get('count', 0)}",
                            f"${readers.get('total_reader_cost', 0):,.2f}"
                        ])
                    else:
                        row.extend(["None", "$0"])
                    
                    row.extend([
                        f"${rec.get('storage_cost', 0):,.2f}",
                        f"${rec.get('total_cost', 0):,.2f}"
                    ])
                    cost_data.append(row)
            
            cost_table = Table(cost_data)
            cost_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1976D2")),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 10),
                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                ('BACKGROUND', (0,1), (-1,-1), colors.HexColor("#E3F2FD")),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('FONTSIZE', (0,1), (-1,-1), 8)
            ]))
            elements.append(cost_table)
            elements.append(Spacer(1, 20))
            
            # Generate cost comparison chart
            cost_chart = self._generate_cost_comparison_chart(calculator)
            if cost_chart:
                elements.append(Image(cost_chart, width=6*inch, height=4*inch))
                elements.append(Spacer(1, 15))
            
            # Page break before detailed recommendations
            elements.append(PageBreak())
            
            # Detailed Environment Recommendations
            elements.append(Paragraph("Detailed Environment Recommendations", styles['Heading1']))
            
            for env in ["PROD", "SQA", "QA", "DEV"]:
                if env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec:
                        elements.append(Paragraph(f"{env} Environment", styles['Heading2']))
                        
                        # Writer details
                        writer = rec.get("writer", {})
                        writer_details = f"""
                        <b>Writer Instance:</b><br/>
                        • Instance Type: {writer.get('instance_type', 'N/A')}<br/>
                        • vCPUs: {writer.get('actual_vCPUs', 0)} (required: {writer.get('vCPUs', 0)})<br/>
                        • RAM: {writer.get('actual_RAM_GB', 0)}GB (required: {writer.get('RAM_GB', 0)}GB)<br/>
                        • Monthly Cost: ${writer.get('monthly_cost', 0):,.2f}<br/>
                        """
                        elements.append(Paragraph(writer_details, styles['Normal']))
                        
                        # Reader details (if applicable)
                        if rec.get("readers"):
                            readers = rec["readers"]
                            reader_details = f"""
                            <b>Reader Instances:</b><br/>
                            • Instance Type: {readers.get('instance_type', 'N/A')} x{readers.get('count', 0)}<br/>
                            • vCPUs per reader: {readers.get('actual_vCPUs', 0)} (required: {readers.get('vCPUs', 0)})<br/>
                            • RAM per reader: {readers.get('actual_RAM_GB', 0)}GB (required: {readers.get('RAM_GB', 0)}GB)<br/>
                            • Total vCPUs: {readers.get('total_vCPUs', 0)}<br/>
                            • Total RAM: {readers.get('total_RAM_GB', 0)}GB<br/>
                            • Total Monthly Cost: ${readers.get('total_reader_cost', 0):,.2f}<br/>
                            """
                            elements.append(Paragraph(reader_details, styles['Normal']))
                        else:
                            elements.append(Paragraph("<b>Reader Instances:</b> None (Single-AZ deployment)", styles['Normal']))
                        
                        # Storage and total cost
                        infra_details = f"""
                        <b>Infrastructure:</b><br/>
                        • Storage: {rec.get('storage_GB', 0)}GB<br/>
                        • Workload Pattern: {rec.get('workload_pattern', 'N/A')}<br/>
                        • Total Monthly Cost: ${rec.get('total_cost', 0):,.2f}<br/>
                        """
                        elements.append(Paragraph(infra_details, styles['Normal']))
                        
                        # Advisories
                        if rec.get("advisories"):
                            elements.append(Paragraph("Optimization Advisories:", styles['Heading3']))
                            for advisory in rec["advisories"]:
                                elements.append(Paragraph(f"• {advisory}", styles['Normal']))
                        
                        elements.append(Spacer(1, 15))
            
            # Resource allocation chart
            resource_chart = self._generate_resource_allocation_chart(calculator)
            if resource_chart:
                elements.append(Paragraph("Resource Allocation Analysis", styles['Heading2']))
                elements.append(Image(resource_chart, width=6*inch, height=4*inch))
                elements.append(Spacer(1, 15))
            
            # Recommendations and Best Practices
            elements.append(PageBreak())
            elements.append(Paragraph("Recommendations and Best Practices", styles['Heading1']))
            
            best_practices = self._generate_best_practices(calculator)
            elements.append(Paragraph(best_practices, styles['Normal']))
            
            # Build PDF
            doc.build(elements)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            return b"PDF generation requires reportlab. Please install it with: pip install reportlab"
    
    def _generate_architecture_description(self, calculator):
        """Generate architecture description based on deployment type"""
        deployment = calculator.inputs.get("deployment", "Multi-AZ")
        deployment_config = calculator.DEPLOYMENT_OPTIONS.get(deployment, {})
        
        description = f"<b>Deployment Configuration:</b> {deployment}<br/>"
        description += f"{deployment_config.get('description', 'Standard deployment')}<br/><br/>"
        
        if deployment_config.get("has_readers", False):
            reader_count = deployment_config.get("reader_count", 0)
            description += f"""
            <b>Architecture Components:</b><br/>
            • <b>Writer Instance:</b> Handles all write operations and approximately 30% of read operations<br/>
            • <b>Reader Instances:</b> {reader_count} instance(s) handle approximately 70% of read operations<br/>
            • <b>Load Distribution:</b> Automatic failover and read traffic distribution across readers<br/>
            • <b>Data Synchronization:</b> Synchronous replication to standby, asynchronous to read replicas<br/>
            """
        else:
            description += """
            <b>Architecture Components:</b><br/>
            • <b>Single Writer Instance:</b> Handles all read and write operations<br/>
            • <b>No Read Replicas:</b> All traffic directed to primary instance<br/>
            • <b>Simplified Management:</b> Single point of management and monitoring<br/>
            """
        
        # Add workload pattern information
        workload_pattern = calculator.inputs.get("workload_pattern", "MIXED")
        if workload_pattern in calculator.WORKLOAD_PATTERNS:
            pattern_info = calculator.WORKLOAD_PATTERNS[workload_pattern]
            description += f"""<br/>
            <b>Workload Characteristics:</b><br/>
            • <b>Read Operations:</b> {pattern_info['read_percentage']}%<br/>
            • <b>Write Operations:</b> {pattern_info['write_percentage']}%<br/>
            • <b>Pattern:</b> {pattern_info['description']}<br/>
            """
        
        return description
    
    def _generate_cost_comparison_chart(self, calculator):
        """Generate cost comparison chart"""
        try:
            plt.style.use('default')
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 6))
            
            # Prepare data
            environments = []
            writer_costs = []
            reader_costs = []
            storage_costs = []
            total_costs = []
            
            for env in ["DEV", "QA", "SQA", "PROD"]:
                if env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec:
                        environments.append(env)
                        writer_costs.append(rec.get('instance_cost', 0))
                        reader_costs.append(rec.get('reader_cost', 0))
                        storage_costs.append(rec.get('storage_cost', 0))
                        total_costs.append(rec.get('total_cost', 0))
            
            if not environments:
                return None
            
            # Chart 1: Stacked cost breakdown
            width = 0.6
            x = np.arange(len(environments))
            
            ax1.bar(x, writer_costs, width, label='Writer', color=self.colors['writer'])
            ax1.bar(x, reader_costs, width, bottom=writer_costs, label='Readers', color=self.colors['reader'])
            
            # Calculate storage bottom
            storage_bottom = [w + r for w, r in zip(writer_costs, reader_costs)]
            ax1.bar(x, storage_costs, width, bottom=storage_bottom, label='Storage', color=self.colors['storage'])
            
            ax1.set_title('Monthly Cost Breakdown by Environment', fontsize=14, fontweight='bold')
            ax1.set_xlabel('Environment')
            ax1.set_ylabel('Monthly Cost ($)')
            ax1.set_xticks(x)
            ax1.set_xticklabels(environments)
            ax1.legend()
            ax1.grid(True, alpha=0.3)
            
            # Add value labels on bars
            for i, total in enumerate(total_costs):
                ax1.text(i, total + max(total_costs) * 0.01, f'${total:,.0f}', 
                        ha='center', va='bottom', fontweight='bold')
            
            # Chart 2: Reader vs Writer cost ratio (only for environments with readers)
            envs_with_readers = []
            writer_percentages = []
            reader_percentages = []
            
            for i, env in enumerate(environments):
                if reader_costs[i] > 0:
                    envs_with_readers.append(env)
                    total_instance_cost = writer_costs[i] + reader_costs[i]
                    writer_pct = (writer_costs[i] / total_instance_cost) * 100
                    reader_pct = (reader_costs[i] / total_instance_cost) * 100
                    writer_percentages.append(writer_pct)
                    reader_percentages.append(reader_pct)
            
            if envs_with_readers:
                x2 = np.arange(len(envs_with_readers))
                ax2.bar(x2, writer_percentages, width, label='Writer %', color=self.colors['writer'])
                ax2.bar(x2, reader_percentages, width, bottom=writer_percentages, 
                       label='Reader %', color=self.colors['reader'])
                
                ax2.set_title('Writer vs Reader Cost Distribution', fontsize=14, fontweight='bold')
                ax2.set_xlabel('Environment')
                ax2.set_ylabel('Percentage of Instance Costs')
                ax2.set_xticks(x2)
                ax2.set_xticklabels(envs_with_readers)
                ax2.legend()
                ax2.grid(True, alpha=0.3)
                ax2.set_ylim(0, 100)
                
                # Add percentage labels
                for i, (w_pct, r_pct) in enumerate(zip(writer_percentages, reader_percentages)):
                    ax2.text(i, w_pct/2, f'{w_pct:.0f}%', ha='center', va='center', 
                            fontweight='bold', color='white')
                    ax2.text(i, w_pct + r_pct/2, f'{r_pct:.0f}%', ha='center', va='center',
                            fontweight='bold', color='white')
            else:
                ax2.text(0.5, 0.5, 'No Multi-AZ Deployments\n(Single-AZ has no readers)', 
                        ha='center', va='center', transform=ax2.transAxes, 
                        fontsize=12, style='italic')
                ax2.set_title('Writer vs Reader Cost Distribution', fontsize=14, fontweight='bold')
            
            plt.tight_layout()
            
            # Save to buffer
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            return buffer
            
        except Exception as e:
            print(f"Error generating cost chart: {e}")
            return None
    
    def _generate_resource_allocation_chart(self, calculator):
        """Generate resource allocation chart"""
        try:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 6))
            
            # Prepare data
            environments = []
            writer_vcpus = []
            reader_vcpus = []
            writer_ram = []
            reader_ram = []
            
            for env in ["DEV", "QA", "SQA", "PROD"]:
                if env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec:
                        environments.append(env)
                        writer_vcpus.append(rec.get('writer', {}).get('actual_vCPUs', 0))
                        writer_ram.append(rec.get('writer', {}).get('actual_RAM_GB', 0))
                        
                        if rec.get('readers'):
                            reader_vcpus.append(rec['readers'].get('total_vCPUs', 0))
                            reader_ram.append(rec['readers'].get('total_RAM_GB', 0))
                        else:
                            reader_vcpus.append(0)
                            reader_ram.append(0)
            
            if not environments:
                return None
            
            width = 0.6
            x = np.arange(len(environments))
            
            # vCPU Chart
            ax1.bar(x, writer_vcpus, width, label='Writer vCPUs', color=self.colors['writer'])
            ax1.bar(x, reader_vcpus, width, bottom=writer_vcpus, label='Reader vCPUs (Total)', 
                   color=self.colors['reader'])
            
            ax1.set_title('vCPU Allocation: Writer vs Readers', fontsize=14, fontweight='bold')
            ax1.set_xlabel('Environment')
            ax1.set_ylabel('vCPUs')
            ax1.set_xticks(x)
            ax1.set_xticklabels(environments)
            ax1.legend()
            ax1.grid(True, alpha=0.3)
            
            # Add total labels
            total_vcpus = [w + r for w, r in zip(writer_vcpus, reader_vcpus)]
            for i, total in enumerate(total_vcpus):
                ax1.text(i, total + max(total_vcpus) * 0.01, f'{total}', 
                        ha='center', va='bottom', fontweight='bold')
            
            # RAM Chart
            ax2.bar(x, writer_ram, width, label='Writer RAM', color=self.colors['writer'])
            ax2.bar(x, reader_ram, width, bottom=writer_ram, label='Reader RAM (Total)', 
                   color=self.colors['reader'])
            
            ax2.set_title('RAM Allocation: Writer vs Readers', fontsize=14, fontweight='bold')
            ax2.set_xlabel('Environment')
            ax2.set_ylabel('RAM (GB)')
            ax2.set_xticks(x)
            ax2.set_xticklabels(environments)
            ax2.legend()
            ax2.grid(True, alpha=0.3)
            
            # Add total labels
            total_ram = [w + r for w, r in zip(writer_ram, reader_ram)]
            for i, total in enumerate(total_ram):
                ax2.text(i, total + max(total_ram) * 0.01, f'{total}GB', 
                        ha='center', va='bottom', fontweight='bold')
            
            plt.tight_layout()
            
            # Save to buffer
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            return buffer
            
        except Exception as e:
            print(f"Error generating resource chart: {e}")
            return None
    
    def _generate_best_practices(self, calculator):
        """Generate best practices recommendations"""
        deployment = calculator.inputs.get("deployment", "Multi-AZ")
        deployment_config = calculator.DEPLOYMENT_OPTIONS.get(deployment, {})
        has_readers = deployment_config.get("has_readers", False)
        
        practices = """
        <b>Reader/Writer Architecture Best Practices:</b><br/><br/>
        """
        
        if has_readers:
            practices += """
            <b>Multi-AZ with Read Replicas:</b><br/>
            • <b>Connection Management:</b> Use separate connection strings for read and write operations<br/>
            • <b>Application Logic:</b> Direct read-only queries to reader endpoints<br/>
            • <b>Monitoring:</b> Monitor replication lag between writer and readers<br/>
            • <b>Scaling:</b> Consider Aurora Auto Scaling for variable read workloads<br/>
            • <b>Failover:</b> Test automatic failover scenarios regularly<br/><br/>
            
            <b>Performance Optimization:</b><br/>
            • <b>Connection Pooling:</b> Implement connection pooling to reduce connection overhead<br/>
            • <b>Read Distribution:</b> Balance read traffic across available reader instances<br/>
            • <b>Caching:</b> Implement application-level caching for frequently accessed data<br/>
            • <b>Query Optimization:</b> Optimize queries to reduce load on writer instance<br/><br/>
            """
        else:
            practices += """
            <b>Single-AZ Deployment:</b><br/>
            • <b>Backup Strategy:</b> Ensure regular automated backups are configured<br/>
            • <b>Monitoring:</b> Implement comprehensive monitoring for single point of failure<br/>
            • <b>Scaling:</b> Plan for vertical scaling as workload grows<br/>
            • <b>High Availability:</b> Consider migrating to Multi-AZ for production workloads<br/><br/>
            """
        
        practices += """
        <b>Cost Optimization:</b><br/>
        • <b>Reserved Instances:</b> Use Reserved Instances for predictable workloads (up to 75% savings)<br/>
        • <b>Right-sizing:</b> Regularly review and adjust instance sizes based on utilization<br/>
        • <b>Storage Optimization:</b> Use appropriate storage types (gp3 vs io1/io2) based on IOPS requirements<br/>
        • <b>Environment Tiering:</b> Use smaller instances for development and testing environments<br/><br/>
        
        <b>Security and Compliance:</b><br/>
        • <b>Encryption:</b> Enable encryption at rest and in transit<br/>
        • <b>Network Security:</b> Use VPC and security groups to control access<br/>
        • <b>Authentication:</b> Implement IAM database authentication where possible<br/>
        • <b>Auditing:</b> Enable Performance Insights and CloudWatch logging<br/>
        """
        
        return practices
    
    def generate_excel_report(self, calculator):
        """Generate comprehensive Excel report with multiple sheets"""
        buffer = BytesIO()
        
        try:
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                # Summary Sheet
                summary_data = []
                for env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec:
                        row = {
                            "Environment": env,
                            "Deployment": calculator.inputs.get("deployment", "N/A"),
                            "Workload Pattern": rec.get("workload_pattern", "N/A"),
                            "Writer Instance": rec.get("writer", {}).get("instance_type", "N/A"),
                            "Writer vCPUs": rec.get("writer", {}).get("actual_vCPUs", 0),
                            "Writer RAM (GB)": rec.get("writer", {}).get("actual_RAM_GB", 0),
                            "Writer Cost": rec.get("instance_cost", 0),
                        }
                        
                        if rec.get("readers"):
                            readers = rec["readers"]
                            row.update({
                                "Reader Instance": readers.get("instance_type", "N/A"),
                                "Reader Count": readers.get("count", 0),
                                "Reader vCPUs (each)": readers.get("actual_vCPUs", 0),
                                "Reader RAM (each)": readers.get("actual_RAM_GB", 0),
                                "Total Reader vCPUs": readers.get("total_vCPUs", 0),
                                "Total Reader RAM": readers.get("total_RAM_GB", 0),
                                "Reader Cost": readers.get("total_reader_cost", 0)
                            })
                        else:
                            row.update({
                                "Reader Instance": "None",
                                "Reader Count": 0,
                                "Reader vCPUs (each)": 0,
                                "Reader RAM (each)": 0,
                                "Total Reader vCPUs": 0,
                                "Total Reader RAM": 0,
                                "Reader Cost": 0
                            })
                        
                        row.update({
                            "Storage (GB)": rec.get("storage_GB", 0),
                            "Storage Cost": rec.get("storage_cost", 0),
                            "Total Cost": rec.get("total_cost", 0)
                        })
                        
                        summary_data.append(row)
                
                if summary_data:
                    pd.DataFrame(summary_data).to_excel(writer, sheet_name="Summary", index=False)
                
                # Detailed Cost Breakdown
                cost_breakdown_data = []
                for env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec and "cost_breakdown" in rec:
                        breakdown = rec["cost_breakdown"]
                        cost_breakdown_data.append({
                            "Environment": env,
                            "Writer Instance Cost": breakdown.get("writer_monthly", 0),
                            "Reader Instance Cost": breakdown.get("reader_monthly", 0),
                            "Storage Cost": breakdown.get("storage_monthly", 0),
                            "Backup Cost": breakdown.get("backup_monthly", 0),
                            "Features Cost": breakdown.get("features_monthly", 0),
                            "Data Transfer Cost": breakdown.get("data_transfer_monthly", 0),
                            "Total Monthly Cost": breakdown.get("total_monthly", 0)
                        })
                
                if cost_breakdown_data:
                    pd.DataFrame(cost_breakdown_data).to_excel(writer, sheet_name="Cost Breakdown", index=False)
                
                # Resource Comparison
                resource_data = []
                for env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec:
                        writer = rec.get("writer", {})
                        readers = rec.get("readers", {})
                        
                        resource_data.append({
                            "Environment": env,
                            "Writer Required vCPUs": writer.get("vCPUs", 0),
                            "Writer Allocated vCPUs": writer.get("actual_vCPUs", 0),
                            "Writer Required RAM": writer.get("RAM_GB", 0),
                            "Writer Allocated RAM": writer.get("actual_RAM_GB", 0),
                            "Reader Required vCPUs (each)": readers.get("vCPUs", 0) if readers else 0,
                            "Reader Allocated vCPUs (each)": readers.get("actual_vCPUs", 0) if readers else 0,
                            "Reader Required RAM (each)": readers.get("RAM_GB", 0) if readers else 0,
                            "Reader Allocated RAM (each)": readers.get("actual_RAM_GB", 0) if readers else 0,
                            "Reader Count": readers.get("count", 0) if readers else 0,
                            "Total Reader vCPUs": readers.get("total_vCPUs", 0) if readers else 0,
                            "Total Reader RAM": readers.get("total_RAM_GB", 0) if readers else 0
                        })
                
                if resource_data:
                    pd.DataFrame(resource_data).to_excel(writer, sheet_name="Resource Allocation", index=False)
                
                # Advisories
                advisory_data = []
                for env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec and rec.get("advisories"):
                        for i, advisory in enumerate(rec["advisories"]):
                            advisory_data.append({
                                "Environment": env,
                                "Advisory #": i + 1,
                                "Advisory": advisory
                            })
                
                if advisory_data:
                    pd.DataFrame(advisory_data).to_excel(writer, sheet_name="Advisories", index=False)
                
                # Configuration Summary
                config_data = [{
                    "Parameter": "Database Engine",
                    "Value": calculator.inputs.get("engine", "N/A")
                }, {
                    "Parameter": "Region", 
                    "Value": calculator.inputs.get("region", "N/A")
                }, {
                    "Parameter": "Deployment Type",
                    "Value": calculator.inputs.get("deployment", "N/A")
                }, {
                    "Parameter": "Workload Pattern",
                    "Value": calculator.inputs.get("workload_pattern", "N/A")
                }, {
                    "Parameter": "Read/Write Ratio",
                    "Value": calculator.inputs.get("read_write_ratio", "N/A")
                }, {
                    "Parameter": "Current CPU Cores",
                    "Value": calculator.inputs.get("on_prem_cores", "N/A")
                }, {
                    "Parameter": "Peak CPU Utilization %",
                    "Value": calculator.inputs.get("peak_cpu_percent", "N/A")
                }, {
                    "Parameter": "Current RAM (GB)",
                    "Value": calculator.inputs.get("on_prem_ram_gb", "N/A")
                }, {
                    "Parameter": "Peak RAM Utilization %",
                    "Value": calculator.inputs.get("peak_ram_percent", "N/A")
                }, {
                    "Parameter": "Current Storage (GB)",
                    "Value": calculator.inputs.get("storage_current_gb", "N/A")
                }]
                
                pd.DataFrame(config_data).to_excel(writer, sheet_name="Configuration", index=False)
        
        except Exception as e:
            print(f"Error generating Excel report: {e}")
            # Fallback to simple CSV
            summary_data = []
            for env in calculator.recommendations:
                rec = calculator.recommendations[env]
                if "error" not in rec:
                    summary_data.append({
                        "Environment": env,
                        "Writer Instance": rec.get("writer", {}).get("instance_type", "N/A"),
                        "Reader Instance": f"{rec['readers']['instance_type']} x{rec['readers']['count']}" if rec.get('readers') else "None",
                        "Total Cost": rec.get("total_cost", 0)
                    })
            
            pd.DataFrame(summary_data).to_csv(buffer, index=False)
        
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_docx_report(self, calculator):
        """Generate enhanced DOCX report with reader/writer details"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            doc = Document()
            
            # Title
            title = doc.add_heading("AWS RDS/Aurora Sizing Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading("Reader/Writer Optimized Recommendations", level=2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Executive Summary
            doc.add_heading("Executive Summary", level=1)
            
            deployment = calculator.inputs.get("deployment", "Multi-AZ")
            deployment_config = calculator.DEPLOYMENT_OPTIONS.get(deployment, {})
            
            summary_p = doc.add_paragraph()
            summary_p.add_run("Migration Sizing Report\n").bold = True
            summary_p.add_run(f"Database Engine: {calculator.inputs.get('engine', 'N/A')}\n")
            summary_p.add_run(f"Region: {calculator.inputs.get('region', 'N/A')}\n")
            summary_p.add_run(f"Deployment Type: {deployment}\n")
            summary_p.add_run(f"Reader Instances: {'Yes (' + str(deployment_config.get('reader_count', 0)) + ' per environment)' if deployment_config.get('has_readers') else 'No (Single-AZ)'}\n")
            summary_p.add_run(f"Workload Pattern: {calculator.inputs.get('workload_pattern', 'N/A')}\n")
            
            if "PROD" in calculator.recommendations:
                prod_rec = calculator.recommendations["PROD"]
                if "error" not in prod_rec:
                    summary_p.add_run(f"Production Writer: {prod_rec.get('writer', {}).get('instance_type', 'N/A')}\n")
                    if prod_rec.get("readers"):
                        readers = prod_rec["readers"]
                        summary_p.add_run(f"Production Readers: {readers.get('instance_type', 'N/A')} x{readers.get('count', 0)}\n")
                    summary_p.add_run(f"Total Monthly Cost: ${prod_rec.get('total_cost', 0):,.2f}\n")
            
            # Detailed Recommendations
            doc.add_heading("Environment-Specific Recommendations", level=1)
            
            # Create comprehensive table
            table = doc.add_table(rows=1, cols=8)
            table.style = "Light Shading Accent 1"
            
            hdr_cells = table.rows[0].cells
            headers = ["Environment", "Writer Instance", "Writer Cost", "Reader Instance", 
                      "Reader Count", "Reader Cost", "Storage (GB)", "Total Cost"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            for env in ["PROD", "SQA", "QA", "DEV"]:
                if env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec:
                        row_cells = table.add_row().cells
                        row_cells[0].text = env
                        row_cells[1].text = rec.get("writer", {}).get("instance_type", "N/A")
                        row_cells[2].text = f"${rec.get('instance_cost', 0):,.2f}"
                        
                        if rec.get("readers"):
                            readers = rec["readers"]
                            row_cells[3].text = readers.get("instance_type", "N/A")
                            row_cells[4].text = str(readers.get("count", 0))
                            row_cells[5].text = f"${readers.get('total_reader_cost', 0):,.2f}"
                        else:
                            row_cells[3].text = "None"
                            row_cells[4].text = "0"
                            row_cells[5].text = "$0"
                        
                        row_cells[6].text = str(rec.get("storage_GB", 0))
                        row_cells[7].text = f"${rec.get('total_cost', 0):,.2f}"
            
            # Detailed Environment Analysis
            doc.add_heading("Detailed Environment Analysis", level=1)
            
            for env in ["PROD", "SQA", "QA", "DEV"]:
                if env in calculator.recommendations:
                    rec = calculator.recommendations[env]
                    if "error" not in rec:
                        doc.add_heading(f"{env} Environment", level=2)
                        
                        # Writer details
                        writer = rec.get("writer", {})
                        writer_p = doc.add_paragraph()
                        writer_p.add_run("Writer Instance:\n").bold = True
                        writer_p.add_run(f"• Instance Type: {writer.get('instance_type', 'N/A')}\n")
                        writer_p.add_run(f"• vCPUs: {writer.get('actual_vCPUs', 0)} (required: {writer.get('vCPUs', 0)})\n")
                        writer_p.add_run(f"• RAM: {writer.get('actual_RAM_GB', 0)}GB (required: {writer.get('RAM_GB', 0)}GB)\n")
                        writer_p.add_run(f"• Monthly Cost: ${writer.get('monthly_cost', 0):,.2f}\n")
                        
                        # Reader details
                        if rec.get("readers"):
                            readers = rec["readers"]
                            reader_p = doc.add_paragraph()
                            reader_p.add_run("Reader Instances:\n").bold = True
                            reader_p.add_run(f"• Instance Type: {readers.get('instance_type', 'N/A')} x{readers.get('count', 0)}\n")
                            reader_p.add_run(f"• vCPUs per reader: {readers.get('actual_vCPUs', 0)} (required: {readers.get('vCPUs', 0)})\n")
                            reader_p.add_run(f"• RAM per reader: {readers.get('actual_RAM_GB', 0)}GB (required: {readers.get('RAM_GB', 0)}GB)\n")
                            reader_p.add_run(f"• Total vCPUs: {readers.get('total_vCPUs', 0)}\n")
                            reader_p.add_run(f"• Total RAM: {readers.get('total_RAM_GB', 0)}GB\n")
                            reader_p.add_run(f"• Total Monthly Cost: ${readers.get('total_reader_cost', 0):,.2f}\n")
                        else:
                            reader_p = doc.add_paragraph()
                            reader_p.add_run("Reader Instances: ").bold = True
                            reader_p.add_run("None (Single-AZ deployment)")
                        
                        # Workload pattern
                        workload_p = doc.add_paragraph()
                        workload_p.add_run("Workload Pattern: ").bold = True
                        workload_p.add_run(rec.get("workload_pattern", "N/A"))
                        
                        # Storage and total
                        summary_p = doc.add_paragraph()
                        summary_p.add_run(f"Storage: {rec.get('storage_GB', 0)}GB\n")
                        summary_p.add_run(f"Total Monthly Cost: ${rec.get('total_cost', 0):,.2f}\n").bold = True
                        
                        # Advisories
                        if rec.get("advisories"):
                            doc.add_heading("Optimization Advisories:", level=3)
                            for advisory in rec["advisories"]:
                                p = doc.add_paragraph(style='ListBullet')
                                p.add_run(advisory.replace('⚠️', '').replace('💡', '').replace('🚨', '').replace('📖', '').replace('🔄', '').replace('🚀', '').strip())
            
            # Best Practices
            doc.add_heading("Recommendations and Best Practices", level=1)
            
            practices_text = self._generate_best_practices(calculator)
            # Convert HTML-like formatting to Word formatting
            practices_lines = practices_text.replace('<b>', '').replace('</b>', '').replace('<br/>', '\n').split('\n')
            
            for line in practices_lines:
                if line.strip():
                    p = doc.add_paragraph()
                    if line.strip().endswith(':'):
                        p.add_run(line.strip()).bold = True
                    else:
                        p.add_run(line.strip())
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            return b"DOCX generation requires python-docx. Please install it with: pip install python-docx"


# Testing the enhanced report generator
if __name__ == "__main__":
    print("🧪 Testing Enhanced Report Generator with Reader/Writer Support")
    
    # Create a mock calculator with reader/writer recommendations
    class MockCalculator:
        def __init__(self):
            self.inputs = {
                "engine": "aurora-postgresql",
                "region": "us-east-1",
                "deployment": "Multi-AZ Cluster",
                "workload_pattern": "READ_HEAVY",
                "read_write_ratio": "80:20",
                "on_prem_cores": 16,
                "peak_cpu_percent": 70,
                "on_prem_ram_gb": 64,
                "peak_ram_percent": 75,
                "storage_current_gb": 500
            }
            
            self.DEPLOYMENT_OPTIONS = {
                'Multi-AZ Cluster': {
                    'cost_multiplier': 2.5,
                    'has_readers': True,
                    'reader_count': 2,
                    'description': 'Primary + 2 readers across AZs'
                }
            }
            
            self.WORKLOAD_PATTERNS = {
                'READ_HEAVY': {
                    'read_percentage': 80,
                    'write_percentage': 20,
                    'description': 'Read-heavy analytical workload'
                }
            }
            
            self.recommendations = {
                "PROD": {
                    "environment": "PROD",
                    "deployment_type": "Multi-AZ Cluster",
                    "workload_pattern": "80% reads, 20% writes",
                    "writer": {
                        "instance_type": "db.r5.2xlarge",
                        "vCPUs": 6,
                        "RAM_GB": 48,
                        "actual_vCPUs": 8,
                        "actual_RAM_GB": 64,
                        "monthly_cost": 2304
                    },
                    "readers": {
                        "instance_type": "db.r5.xlarge",
                        "vCPUs": 3,
                        "RAM_GB": 24,
                        "actual_vCPUs": 4,
                        "actual_RAM_GB": 32,
                        "count": 2,
                        "total_vCPUs": 8,
                        "total_RAM_GB": 64,
                        "total_reader_cost": 2304
                    },
                    "storage_GB": 650,
                    "instance_cost": 2304,
                    "reader_cost": 2304,
                    "storage_cost": 65,
                    "total_cost": 4733,
                    "advisories": [
                        "🚀 Consider Aurora Auto Scaling for readers based on CPU utilization",
                        "📖 High read workload optimally distributed across readers"
                    ],
                    "cost_breakdown": {
                        "writer_monthly": 2304,
                        "reader_monthly": 2304,
                        "storage_monthly": 65,
                        "backup_monthly": 60,
                        "total_monthly": 4733
                    }
                },
                "DEV": {
                    "environment": "DEV",
                    "deployment_type": "Multi-AZ Cluster",
                    "workload_pattern": "80% reads, 20% writes",
                    "writer": {
                        "instance_type": "db.t3.large",
                        "vCPUs": 2,
                        "RAM_GB": 6,
                        "actual_vCPUs": 2,
                        "actual_RAM_GB": 8,
                        "monthly_cost": 147
                    },
                    "readers": {
                        "instance_type": "db.t3.medium",
                        "vCPUs": 1,
                        "RAM_GB": 3,
                        "actual_vCPUs": 2,
                        "actual_RAM_GB": 4,
                        "count": 2,
                        "total_vCPUs": 4,
                        "total_RAM_GB": 8,
                        "total_reader_cost": 147
                    },
                    "storage_GB": 195,
                    "instance_cost": 147,
                    "reader_cost": 147,
                    "storage_cost": 20,
                    "total_cost": 334,
                    "advisories": [
                        "💡 Consider smaller instances for development environment",
                        "💰 Single read replica may be sufficient for DEV environment"
                    ],
                    "cost_breakdown": {
                        "writer_monthly": 147,
                        "reader_monthly": 147,
                        "storage_monthly": 20,
                        "backup_monthly": 5,
                        "total_monthly": 334
                    }
                }
            }
    
    # Test the enhanced report generator
    mock_calculator = MockCalculator()
    report_generator = EnhancedReportGenerator()
    
    print("\n📊 Testing Excel Report Generation...")
    try:
        excel_data = report_generator.generate_excel_report(mock_calculator)
        print(f"✅ Excel report generated successfully ({len(excel_data)} bytes)")
    except Exception as e:
        print(f"❌ Excel report generation failed: {e}")
    
    print("\n📄 Testing DOCX Report Generation...")
    try:
        docx_data = report_generator.generate_docx_report(mock_calculator)
        print(f"✅ DOCX report generated successfully ({len(docx_data)} bytes)")
    except Exception as e:
        print(f"❌ DOCX report generation failed: {e}")
    
    print("\n📋 Testing PDF Report Generation...")
    try:
        pdf_data = report_generator.generate_pdf_report(mock_calculator)
        print(f"✅ PDF report generated successfully ({len(pdf_data)} bytes)")
    except Exception as e:
        print(f"❌ PDF report generation failed: {e}")
    
    print(f"\n🎯 Enhanced Report Features:")
    print("✅ Reader/Writer cost breakdown and analysis")
    print("✅ Resource allocation charts (vCPU and RAM distribution)")
    print("✅ Deployment architecture descriptions")
    print("✅ Multi-sheet Excel reports with detailed breakdowns")
    print("✅ Environment-specific recommendations and advisories")
    print("✅ Best practices for reader/writer architectures")
    print("✅ Visual cost comparison charts")
    print("✅ Comprehensive configuration summaries")